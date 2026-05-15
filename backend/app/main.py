from app.api import v3_evidence
from app.api import v3_data_quality
from app.api import v3_analysis
import csv
import io
import os
import uuid
from datetime import datetime
from decimal import Decimal
from typing import Any, Optional

import psycopg2
import psycopg2.extras
from fastapi import FastAPI, File, Form, Query, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse

DATABASE_URL = os.getenv(
    "DATABASE_URL",
    "postgresql://shield_user:shield_pass_123@127.0.0.1:5432/shield_monitor",
)
DEFAULT_SECTION_ID = os.getenv("DEFAULT_SECTION_ID", "33333333-3333-3333-3333-333333333333")
CURRENT_RING_NO = int(os.getenv("CURRENT_RING_NO", "336"))

app = FastAPI(title="Shield Monitor Platform V2", version="2.3.0")
from app.api import v4_engineering
from app.api import ai_diagnosis
from app.api import v4_position_context
from app.api import v4_nearby_monitoring
from app.api import v4_intelligent_analysis
from app.api import v4_report_cockpit
from app.api import v4_specialized_pages_v2
app.include_router(v4_engineering.router)
app.include_router(v4_position_context.router)
app.include_router(v4_nearby_monitoring.router)
app.include_router(ai_diagnosis.router)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)


def conn():
    return psycopg2.connect(DATABASE_URL)


def one(sql: str, params: tuple[Any, ...] = ()):  # noqa: ANN401
    with conn() as c:
        with c.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
            cur.execute(sql, params)
            row = cur.fetchone()
            return dict(row) if row else None


def all_rows(sql: str, params: tuple[Any, ...] = ()):  # noqa: ANN401
    with conn() as c:
        with c.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
            cur.execute(sql, params)
            return [dict(r) for r in cur.fetchall()]


def exec_sql(sql: str, params: tuple[Any, ...] = ()):  # noqa: ANN401
    with conn() as c:
        with c.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
            cur.execute(sql, params)
            try:
                row = cur.fetchone()
                return dict(row) if row else None
            except psycopg2.ProgrammingError:
                return None


def num(v):
    if v is None:
        return None
    if isinstance(v, Decimal):
        return float(v)
    return v


def iso(v):
    if v is None:
        return None
    if hasattr(v, "isoformat"):
        return v.isoformat()
    return v


def parse_mileage(value: str | None):
    if not value:
        return None
    text = str(value).strip().upper().replace("DK", "")
    if "+" not in text:
        try:
            return float(text)
        except Exception:
            return None
    km, meter = text.split("+", 1)
    try:
        return float(km) * 1000 + float(meter)
    except Exception:
        return None


def project_summary(section_id: str = DEFAULT_SECTION_ID):
    row = one(
        """
        SELECT p.project_id, p.project_name, p.contractor_name,
               ts.section_id, ts.section_name, ts.start_mileage, ts.end_mileage,
               ts.start_mileage_m, ts.end_mileage_m, ts.length_m, ts.tunnel_form,
               ts.design_speed_kmh, ts.max_burial_depth_m
        FROM project p
        JOIN tunnel_section ts ON ts.project_id = p.project_id
        WHERE ts.section_id = %s::uuid
        LIMIT 1
        """,
        (section_id,),
    )
    if not row:
        return None
    return {
        "projectId": str(row["project_id"]),
        "projectName": row["project_name"],
        "contractorName": row["contractor_name"],
        "sectionId": str(row["section_id"]),
        "sectionName": row["section_name"],
        "startMileage": row["start_mileage"],
        "endMileage": row["end_mileage"],
        "startMileageM": num(row["start_mileage_m"]),
        "endMileageM": num(row["end_mileage_m"]),
        "lengthM": num(row["length_m"]),
        "tunnelForm": row["tunnel_form"],
        "designSpeedKmh": num(row["design_speed_kmh"]),
        "maxBurialDepthM": num(row["max_burial_depth_m"]),
    }


def ring_api(row):
    if not row:
        return None
    return {
        "ringId": str(row["ring_id"]),
        "sectionId": str(row["section_id"]),
        "ringNo": row["ring_no"],
        "workDate": iso(row["work_date"]),
        "startMileage": row["start_mileage"],
        "endMileage": row["end_mileage"],
        "startMileageM": num(row["start_mileage_m"]),
        "endMileageM": num(row["end_mileage_m"]),
        "constructionStage": row["construction_stage"],
        "isActual": row["is_actual"],
    }


def ring_by_no(ring_no: int, section_id: str = DEFAULT_SECTION_ID):
    return ring_api(one("SELECT * FROM ring_mileage_map WHERE section_id=%s::uuid AND ring_no=%s", (section_id, ring_no)))


def current_ring(section_id: str = DEFAULT_SECTION_ID):
    row = one("SELECT * FROM ring_mileage_map WHERE section_id=%s::uuid AND ring_no=%s", (section_id, CURRENT_RING_NO))
    if not row:
        row = one("SELECT * FROM ring_mileage_map WHERE section_id=%s::uuid ORDER BY ring_no DESC LIMIT 1", (section_id,))
    return ring_api(row)


def ring_bounds(section_id: str = DEFAULT_SECTION_ID):
    row = one("SELECT MIN(ring_no) AS min_ring, MAX(ring_no) AS max_ring FROM ring_mileage_map WHERE section_id=%s::uuid", (section_id,))
    return {"minRing": row["min_ring"] or 1, "maxRing": row["max_ring"] or 1}


def risk_api(row, current_mileage_m=None):
    start_m = num(row.get("start_mileage_m"))
    end_m = num(row.get("end_mileage_m"))
    status = "normal"
    distance_to_start = None
    distance_to_end = None
    if current_mileage_m is not None and start_m is not None and end_m is not None:
        distance_to_start = round(start_m - current_mileage_m, 2)
        distance_to_end = round(end_m - current_mileage_m, 2)
        if start_m <= current_mileage_m <= end_m:
            status = "inside"
        elif 0 <= start_m - current_mileage_m <= 120:
            status = "approaching"
        elif current_mileage_m > end_m:
            status = "passed"
    alert = "normal"
    if status == "inside" and row.get("risk_level") == "high":
        alert = "alarm"
    elif status in ("approaching", "inside"):
        alert = "warning"
    return {
        "riskSourceId": str(row["risk_source_id"]),
        "riskName": row["risk_name"],
        "riskType": row["risk_type"],
        "crossingRelation": row["crossing_relation"],
        "startMileage": row["start_mileage"],
        "endMileage": row["end_mileage"],
        "startMileageM": start_m,
        "endMileageM": end_m,
        "minHorizontalDistanceM": num(row.get("min_horizontal_distance_m")),
        "minVerticalDistanceM": num(row.get("min_vertical_distance_m")),
        "protectionLevel": row.get("protection_level"),
        "riskLevel": row.get("risk_level") or "medium",
        "monitoringPointCount": row.get("monitoring_point_count") or 0,
        "status": status,
        "alertLevel": alert,
        "distanceToStartM": distance_to_start,
        "distanceToEndM": distance_to_end,
    }


def risks(section_id: str = DEFAULT_SECTION_ID, ring_no: Optional[int] = None):
    ring = ring_by_no(ring_no, section_id) if ring_no else current_ring(section_id)
    mileage = ring["endMileageM"] if ring else None
    rows = all_rows(
        """
        SELECT rs.*, COUNT(mp.point_id) AS monitoring_point_count
        FROM risk_source rs
        LEFT JOIN monitoring_point mp ON mp.risk_source_id = rs.risk_source_id
        WHERE rs.section_id=%s::uuid
        GROUP BY rs.risk_source_id
        ORDER BY rs.start_mileage_m
        """,
        (section_id,),
    )
    return [risk_api(r, mileage) for r in rows]


def operation_api(row):
    if not row:
        return None
    return {
        "operationId": str(row["operation_id"]),
        "ringNo": row["ring_no"],
        "recordedAt": iso(row["recorded_at"]),
        "advanceSpeed": num(row["advance_speed"]),
        "facePressure": num(row["face_pressure"]),
        "totalThrust": num(row["total_thrust"]),
        "cutterTorque": num(row["cutter_torque"]),
        "cutterRotationSpeed": num(row["cutter_rotation_speed"]),
        "penetration": num(row["penetration"]),
        "slurryInFlow": num(row.get("slurry_in_flow")),
        "slurryOutFlow": num(row.get("slurry_out_flow")),
        "slurryInDensity": num(row.get("slurry_in_density")),
        "slurryOutDensity": num(row.get("slurry_out_density")),
        "alertLevel": row.get("alert_level") or "normal",
    }


def operations(section_id: str = DEFAULT_SECTION_ID, start_ring: int = 320, end_ring: int = 392):
    rows = all_rows(
        """
        SELECT * FROM shield_ring_operation
        WHERE section_id=%s::uuid AND ring_no BETWEEN %s AND %s
        ORDER BY ring_no
        """,
        (section_id, start_ring, end_ring),
    )
    return [operation_api(r) for r in rows]


def operation_for_ring(ring_no: int, section_id: str = DEFAULT_SECTION_ID):
    row = one("SELECT * FROM shield_ring_operation WHERE section_id=%s::uuid AND ring_no=%s LIMIT 1", (section_id, ring_no))
    if not row:
        row = one("SELECT * FROM shield_ring_operation WHERE section_id=%s::uuid ORDER BY recorded_at DESC LIMIT 1", (section_id,))
    return operation_api(row)


def point_api(row):
    return {
        "pointId": str(row["point_id"]),
        "riskSourceId": str(row["risk_source_id"]) if row.get("risk_source_id") else None,
        "riskName": row.get("risk_name"),
        "pointCode": row["point_code"],
        "pointName": row.get("point_name"),
        "monitoringObject": row.get("monitoring_object"),
        "monitoringItem": row["monitoring_item"],
        "mileage": row.get("mileage"),
        "mileageM": num(row.get("mileage_m")),
        "relativePosition": row.get("relative_position"),
        "initialValue": num(row.get("initial_value")),
        "unit": row.get("unit"),
        "warningThreshold": num(row.get("warning_threshold")),
        "alarmThreshold": num(row.get("alarm_threshold")),
        "alertLevel": row.get("alert_level") or "unknown",
    }


def monitoring_points(section_id: str = DEFAULT_SECTION_ID):
    rows = all_rows(
        """
        SELECT mp.*, rs.risk_name, latest.alert_level
        FROM monitoring_point mp
        LEFT JOIN risk_source rs ON rs.risk_source_id=mp.risk_source_id
        LEFT JOIN LATERAL (
            SELECT alert_level FROM monitoring_reading mr
            WHERE mr.point_id=mp.point_id
            ORDER BY measured_at DESC LIMIT 1
        ) latest ON TRUE
        WHERE mp.section_id=%s::uuid
        ORDER BY mp.point_code
        """,
        (section_id,),
    )
    return [point_api(r) for r in rows]


def readings(point_code: str | None = None, point_id: str | None = None):
    if point_id:
        point = one("SELECT mp.*, rs.risk_name FROM monitoring_point mp LEFT JOIN risk_source rs ON rs.risk_source_id=mp.risk_source_id WHERE mp.point_id=%s::uuid", (point_id,))
    elif point_code:
        point = one("SELECT mp.*, rs.risk_name FROM monitoring_point mp LEFT JOIN risk_source rs ON rs.risk_source_id=mp.risk_source_id WHERE mp.point_code=%s", (point_code,))
    else:
        point = one("SELECT mp.*, rs.risk_name FROM monitoring_point mp LEFT JOIN risk_source rs ON rs.risk_source_id=mp.risk_source_id ORDER BY mp.point_code LIMIT 1")
    if not point:
        return {"point": None, "readings": []}
    rows = all_rows("SELECT * FROM monitoring_reading WHERE point_id=%s::uuid ORDER BY measured_at", (str(point["point_id"]),))
    data = [{
        "readingId": str(r["reading_id"]),
        "pointId": str(r["point_id"]),
        "measuredAt": iso(r["measured_at"]),
        "currentValue": num(r["current_value"]),
        "cumulativeChange": num(r["cumulative_change"]),
        "changeRate": num(r["change_rate"]),
        "alertLevel": r["alert_level"],
    } for r in rows]
    return {"point": point_api({**point, "alert_level": data[-1]["alertLevel"] if data else "unknown"}), "readings": data}


def monitoring_summary(section_id: str = DEFAULT_SECTION_ID):
    row = one(
        """
        WITH latest AS (
            SELECT DISTINCT ON (mr.point_id) mr.point_id, mr.cumulative_change, mr.alert_level
            FROM monitoring_reading mr
            JOIN monitoring_point mp ON mp.point_id=mr.point_id
            WHERE mp.section_id=%s::uuid
            ORDER BY mr.point_id, mr.measured_at DESC
        )
        SELECT COUNT(mp.point_id) AS point_count,
               COUNT(*) FILTER (WHERE latest.alert_level='warning') AS warning_count,
               COUNT(*) FILTER (WHERE latest.alert_level='alarm') AS alarm_count,
               MIN(latest.cumulative_change) AS max_settlement
        FROM monitoring_point mp
        LEFT JOIN latest ON latest.point_id=mp.point_id
        WHERE mp.section_id=%s::uuid
        """,
        (section_id, section_id),
    )
    return {
        "pointCount": row.get("point_count") or 0,
        "warningCount": row.get("warning_count") or 0,
        "alarmCount": row.get("alarm_count") or 0,
        "maxSettlement": num(row.get("max_settlement")) or 0,
    }


def event_api(row):
    return {
        "eventId": str(row["event_id"]),
        "ringNo": row.get("ring_no"),
        "riskName": row.get("risk_name"),
        "eventTime": iso(row["event_time"]),
        "eventType": row["event_type"],
        "severity": row["severity"],
        "description": row.get("description"),
        "possibleCause": row.get("possible_cause"),
        "handlingAction": row.get("handling_action"),
        "isShutdown": row.get("is_shutdown"),
        "closureResult": row.get("closure_result"),
        "responsibleParty": row.get("responsible_party"),
    }


def events(section_id: str = DEFAULT_SECTION_ID, ring_no: Optional[int] = None, limit: int = 20):
    if ring_no is None:
        rows = all_rows(
            """
            SELECT e.*, r.ring_no, rs.risk_name
            FROM event_log e
            LEFT JOIN ring_mileage_map r ON r.ring_id=e.ring_id
            LEFT JOIN risk_source rs ON rs.risk_source_id=e.risk_source_id
            WHERE e.section_id=%s::uuid
            ORDER BY e.event_time DESC
            LIMIT %s
            """,
            (section_id, limit),
        )
    else:
        rows = all_rows(
            """
            SELECT e.*, r.ring_no, rs.risk_name, ABS(r.ring_no - %s) AS ring_distance
            FROM event_log e
            LEFT JOIN ring_mileage_map r ON r.ring_id=e.ring_id
            LEFT JOIN risk_source rs ON rs.risk_source_id=e.risk_source_id
            WHERE e.section_id=%s::uuid
            ORDER BY ABS(r.ring_no - %s), e.event_time DESC
            LIMIT %s
            """,
            (ring_no, section_id, ring_no, limit),
        )
    return [event_api(r) for r in rows]


def slurry_grouting(section_id: str = DEFAULT_SECTION_ID, start_ring: int = 320, end_ring: int = 392):
    slurry = all_rows("SELECT * FROM slurry_record WHERE section_id=%s::uuid AND ring_no BETWEEN %s AND %s ORDER BY ring_no", (section_id, start_ring, end_ring))
    grout = all_rows("SELECT * FROM grouting_record WHERE section_id=%s::uuid AND ring_no BETWEEN %s AND %s ORDER BY ring_no", (section_id, start_ring, end_ring))
    grout_by_ring = {g["ring_no"]: g for g in grout}
    items = []
    for s in slurry:
        g = grout_by_ring.get(s["ring_no"], {})
        items.append({
            "ringNo": s["ring_no"],
            "recordedAt": iso(s.get("recorded_at")),
            "slurryInDensity": num(s.get("slurry_in_density")),
            "slurryOutDensity": num(s.get("slurry_out_density")),
            "viscosity": num(s.get("viscosity")),
            "sandContent": num(s.get("sand_content")),
            "phValue": num(s.get("ph_value")),
            "waterLoss": num(s.get("water_loss")),
            "groutingVolume": num(g.get("grouting_volume")),
            "groutingPressure": num(g.get("grouting_pressure")),
            "materialRatio": g.get("material_ratio"),
            "isSecondaryGrouting": g.get("is_secondary_grouting"),
        })
    return items


def table_count(table: str):
    try:
        row = one(f"SELECT COUNT(*) AS count FROM {table}")
        return int(row["count"] or 0)
    except Exception:
        return 0


def system_status():
    names = [
        "source_document", "import_batch", "import_raw_row", "field_mapping",
        "project", "tunnel_section", "ring_mileage_map", "risk_source",
        "monitoring_point", "monitoring_reading", "shield_ring_operation",
        "slurry_record", "grouting_record", "event_log"
    ]
    counts = [{"tableName": n, "rowCount": table_count(n)} for n in names]
    lookup = {x["tableName"]: x["rowCount"] for x in counts}
    checks = [
        {"key": "ring", "name": "环号-里程-日期", "count": lookup.get("ring_mileage_map", 0), "ok": lookup.get("ring_mileage_map", 0) > 0},
        {"key": "risk", "name": "风险源台账", "count": lookup.get("risk_source", 0), "ok": lookup.get("risk_source", 0) > 0},
        {"key": "operation", "name": "盾构掘进参数", "count": lookup.get("shield_ring_operation", 0), "ok": lookup.get("shield_ring_operation", 0) > 0},
        {"key": "monitoring", "name": "监测点与监测日报", "count": lookup.get("monitoring_reading", 0), "ok": lookup.get("monitoring_point", 0) > 0 and lookup.get("monitoring_reading", 0) > 0},
        {"key": "event", "name": "事件与处置闭环", "count": lookup.get("event_log", 0), "ok": lookup.get("event_log", 0) > 0},
    ]
    score = round(sum(1 for c in checks if c["ok"]) / len(checks) * 100)
    return {"mode": "database-postgresql", "tableCounts": counts, "dataQuality": {"readyScore": score, "checks": checks}}


def dashboard(section_id: str = DEFAULT_SECTION_ID, ring_no: Optional[int] = None):
    project = project_summary(section_id)
    current = current_ring(section_id)
    selected = ring_by_no(ring_no, section_id) if ring_no else current
    effective_no = selected["ringNo"] if selected else CURRENT_RING_NO
    selected_mileage = selected["endMileageM"] if selected else None
    bounds = ring_bounds(section_id)
    risk_items = risks(section_id, effective_no)
    active = [r for r in risk_items if r["status"] in ("approaching", "inside")]
    op = operation_for_ring(effective_no, section_id)
    trend = operations(section_id, max(bounds["minRing"], effective_no - 45), min(bounds["maxRing"], effective_no + 45))
    recent = events(section_id, effective_no, 6)
    progress = 0
    if project and selected and project.get("startMileageM") and project.get("endMileageM"):
        progress = round((selected_mileage - project["startMileageM"]) / (project["endMileageM"] - project["startMileageM"]) * 100, 2)
    return {
        "project": project,
        "ringBounds": bounds,
        "currentRing": {
            "ringId": current["ringId"], "ringNo": current["ringNo"], "mileage": current["endMileage"],
            "stage": current["constructionStage"], "progressPercent": progress if current["ringNo"] == effective_no else None,
            "workDate": current["workDate"],
        } if current else None,
        "selectedRing": {
            "ringId": selected["ringId"], "ringNo": selected["ringNo"], "mileage": selected["endMileage"],
            "stage": selected["constructionStage"], "progressPercent": progress, "workDate": selected["workDate"],
        } if selected else None,
        "activeRiskSources": active,
        "allRiskSources": risk_items,
        "operationSummary": op,
        "operationTrend": trend,
        "monitoringSummary": monitoring_summary(section_id),
        "recentEvents": recent,
        "dataUpdatedAt": datetime.now().isoformat(),
    }


ALIASES = {
    "ring_no": ["环号", "施工环", "ring", "ring_no", "环"],
    "recorded_at": ["时间", "日期", "时间戳", "施工日期"],
    "advance_speed": ["推进速度", "掘进速度", "速度"],
    "face_pressure": ["切口压力", "掌子面压力", "仓压", "泥水压力"],
    "total_thrust": ["总推力", "推进力", "推力"],
    "cutter_torque": ["刀盘扭矩", "扭矩"],
    "cutter_rotation_speed": ["刀盘转速", "转速"],
    "point_code": ["测点编号", "点号", "测点"],
    "cumulative_change": ["累计变化", "累计沉降", "累计位移"],
    "change_rate": ["变化速率", "速率"],
}


def suggest(headers):
    items = []
    for h in headers:
        lower = str(h).strip().lower()
        match = None
        confidence = 0
        for key, arr in ALIASES.items():
            for alias in arr:
                if lower == alias.lower() or alias.lower() in lower or lower in alias.lower():
                    match = key
                    confidence = 0.95 if lower == alias.lower() else 0.78
                    break
            if match:
                break
        items.append({
            "sourceFieldName": h,
            "suggestedStandardField": match,
            "confidence": confidence,
            "required": match in ("ring_no", "point_code", "recorded_at"),
            "status": "matched" if confidence >= 0.9 else ("need_confirm" if match else "missing"),
        })
    return items


def parse_upload(filename: str, content: bytes):
    suffix = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    rows: list[dict[str, Any]] = []
    headers: list[str] = []
    if suffix == "csv":
        text = content.decode("utf-8-sig", errors="ignore")
        reader = csv.DictReader(io.StringIO(text))
        headers = list(reader.fieldnames or [])
        rows = [dict(r) for r in reader]
    elif suffix in ("xlsx", "xlsm"):
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(content), data_only=True)
        ws = wb.active
        raw_headers = [c.value for c in next(ws.iter_rows(min_row=1, max_row=1))]
        headers = [str(h).strip() for h in raw_headers if h is not None]
        for row in ws.iter_rows(min_row=2, values_only=True):
            item = {headers[i]: row[i] for i in range(min(len(headers), len(row)))}
            if any(v is not None and str(v).strip() != "" for v in item.values()):
                rows.append(item)
    elif suffix == "docx":
        from docx import Document
        doc = Document(io.BytesIO(content))
        if doc.tables:
            table = doc.tables[0]
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            for tr in table.rows[1:]:
                vals = [cell.text.strip() for cell in tr.cells]
                rows.append({headers[i]: vals[i] for i in range(min(len(headers), len(vals)))})
    else:
        raise ValueError("Only csv/xlsx/docx are supported")
    return headers, rows[:200]


@app.get("/api/v2/health")
def health():
    return {"status": "ok", "mode": "database-postgresql", "version": "2.3.0"}


@app.get("/api/v2/dashboard/overview")
def api_dashboard(section_id: str = Query(DEFAULT_SECTION_ID), ring_no: Optional[int] = Query(None)):
    return dashboard(section_id, ring_no)


@app.get("/api/v2/rings/current")
def api_current_ring(section_id: str = Query(DEFAULT_SECTION_ID)):
    return current_ring(section_id)


@app.get("/api/v2/rings/timeline")
def api_ring_timeline(section_id: str = Query(DEFAULT_SECTION_ID), start_ring: int = 1, end_ring: int = 392):
    rows = all_rows("SELECT * FROM ring_mileage_map WHERE section_id=%s::uuid AND ring_no BETWEEN %s AND %s ORDER BY ring_no", (section_id, start_ring, end_ring))
    return {"items": [ring_api(r) for r in rows]}


@app.get("/api/v2/risk-sources")
def api_risks(section_id: str = Query(DEFAULT_SECTION_ID), ring_no: Optional[int] = Query(None)):
    return {"items": risks(section_id, ring_no)}


@app.get("/api/v2/monitoring/points")
def api_points(section_id: str = Query(DEFAULT_SECTION_ID)):
    return {"items": monitoring_points(section_id)}


@app.get("/api/v2/monitoring/readings")
def api_readings(point_code: Optional[str] = None, point_id: Optional[str] = None):
    return readings(point_code, point_id)


@app.get("/api/v2/shield/ring-operations")
def api_operations(section_id: str = Query(DEFAULT_SECTION_ID), start_ring: int = 320, end_ring: int = 392):
    return {"items": operations(section_id, start_ring, end_ring)}


@app.get("/api/v2/slurry-grouting/records")
def api_slurry(section_id: str = Query(DEFAULT_SECTION_ID), start_ring: int = 320, end_ring: int = 392):
    return {"items": slurry_grouting(section_id, start_ring, end_ring)}


@app.get("/api/v2/events")
def api_events(section_id: str = Query(DEFAULT_SECTION_ID), ring_no: Optional[int] = None, limit: int = 20):
    return {"items": events(section_id, ring_no, limit)}


@app.get("/api/v2/system/status")
def api_status():
    return system_status()


@app.get("/api/v2/sources")
def api_sources():
    rows = all_rows("SELECT * FROM source_document ORDER BY created_at DESC LIMIT 50")
    return {"items": [{"sourceId": str(r["source_id"]), "fileName": r["file_name"], "fileType": r["file_type"], "documentDate": iso(r.get("document_date")), "description": r.get("description")} for r in rows]}


@app.post("/api/v2/imports/upload")
async def upload_import(file: UploadFile = File(...), data_category: str = Form("shield_operation")):
    content = await file.read()
    headers, rows = parse_upload(file.filename or "upload", content)
    source_id = str(uuid.uuid4())
    batch_id = str(uuid.uuid4())
    try:
        exec_sql(
            """
            INSERT INTO source_document (source_id, file_name, file_type, source_department, owner_name, version_label, document_date, description)
            VALUES (%s::uuid, %s, %s, %s, %s, %s, CURRENT_DATE, %s)
            """,
            (source_id, file.filename, file.filename.rsplit('.', 1)[-1] if file.filename and '.' in file.filename else 'unknown', 'upload', 'operator', 'v2-upload', 'V2 data intake upload'),
        )
        exec_sql(
            """
            INSERT INTO import_batch (batch_id, source_id, data_category, original_sheet_name, header_row_index, status)
            VALUES (%s::uuid, %s::uuid, %s, %s, 1, 'uploaded')
            """,
            (batch_id, source_id, data_category, 'default'),
        )
        for i, row in enumerate(rows, 1):
            exec_sql(
                "INSERT INTO import_raw_row (batch_id, row_no, raw_data, validation_status) VALUES (%s::uuid, %s, %s, 'pending')",
                (batch_id, i, psycopg2.extras.Json(row)),
            )
    except Exception as e:
        return JSONResponse(status_code=200, content={"batchId": batch_id, "headers": headers, "sampleRows": rows[:5], "mappingSuggestions": suggest(headers), "warning": str(e)})
    return {"batchId": batch_id, "sourceId": source_id, "dataCategory": data_category, "headers": headers, "rowCount": len(rows), "sampleRows": rows[:5], "mappingSuggestions": suggest(headers), "status": "uploaded"}


@app.post("/api/v2/imports/{batch_id}/mapping")
def save_mapping(batch_id: str, payload: dict[str, Any]):
    return {"batchId": batch_id, "status": "mapping_saved", "payload": payload}


@app.post("/api/v2/imports/{batch_id}/validate")
def validate_import(batch_id: str):
    row = one("SELECT COUNT(*) AS count FROM import_raw_row WHERE batch_id=%s::uuid", (batch_id,))
    total = int(row["count"] or 0) if row else 0
    return {"batchId": batch_id, "status": "validated", "totalRows": total, "validRows": total, "invalidRows": 0, "errors": []}


@app.post("/api/v2/imports/{batch_id}/commit")
def commit_import(batch_id: str):
    row = one("SELECT COUNT(*) AS count FROM import_raw_row WHERE batch_id=%s::uuid", (batch_id,))
    total = int(row["count"] or 0) if row else 0
    exec_sql("UPDATE import_batch SET status='committed', committed_at=now() WHERE batch_id=%s::uuid", (batch_id,))
    return {"batchId": batch_id, "status": "committed", "insertedRows": 0, "rawRowsStored": total, "message": "Raw rows stored. Standard-table commit can be enabled per data category."}

app.include_router(v3_analysis.router, prefix="/api/v3/analysis", tags=["v3-analysis"])

app.include_router(v3_evidence.router)

app.include_router(v3_data_quality.router, prefix="/api/v3", tags=["v3-data-quality"])

app.include_router(v4_intelligent_analysis.router)
app.include_router(v4_report_cockpit.router)
app.include_router(v4_specialized_pages_v2.router)
# V3.11: TBM legacy realtime API adapter
try:
    from app.api import tbm as tbm_router
    app.include_router(tbm_router.router, prefix="/api/tbm", tags=["tbm"])
except Exception as exc:
    print(f"[WARN] TBM router not loaded: {exc}")

# V3.11.1: TBM compatibility route for older frontend calls
try:
    if tbm_router is not None:
        app.include_router(tbm_router.router, prefix="/tbm", tags=["tbm-compat"])
except Exception as exc:
    print(f"[WARN] TBM compatibility router not loaded: {exc}")

# V3.14: PostgreSQL historical-file API routes
try:
    from app.api import file_monitoring as file_monitoring_router
    for _route in reversed(file_monitoring_router.router.routes):
        app.router.routes.insert(0, _route)
except Exception as exc:
    print(f"[WARN] file monitoring router not loaded: {exc}")

# ===== V4_1_BACKEND_ENGINEERING_APIS_BEGIN =====
# Backend-only engineering APIs.
# Purpose:
# - Keep frontend API entry on 8100.
# - Keep 19090 as TBM collector/latest-view source only.
# - Query PostgreSQL with existing conn()/one()/all_rows() helpers already defined above.
import json as _v41_json
import os as _v41_os
import urllib.parse as _v41_urlparse
import urllib.request as _v41_urlrequest
from collections import deque as _v41_deque
from pathlib import Path as _v41_Path


_V41_TBM_BASE = _v41_os.getenv("TBM_API_BASE", "http://127.0.0.1:19090").rstrip("/")
_V41_TBM_HISTORY_JSONL = _v41_Path(_v41_os.getenv("TBM_HISTORY_JSONL", "/opt/tbm_receiver/history.jsonl"))

_V41_TBM_KEYS = [
    "currentRing", "hydraulicOilTemp",
    "cutterStatus", "cutterSpeed", "cutterAngle", "cutterTorque",
    "advanceStatus", "advancePumpPressure", "advanceSpeed", "penetration", "totalThrust", "advanceSpeedSet",
    "chamberPressure1", "chamberPressure2", "chamberPressure3",
    "slurryOutDensity", "slurryOutFlow", "slurryInDensity", "slurryInFlow", "slurryInPressure",
    "shieldTailGap1", "shieldTailGap2", "shieldTailGap3",
    "propelPressureA", "propelPressureB", "propelPressureC", "propelPressureD", "propelPressureE", "propelPressureF",
    "groutTotal", "segmentPosition",
]


def _v41_ok(data, source="engineering"):
    return {"code": 0, "data": data, "source": source}


def _v41_limit(value, default=100, maximum=2000):
    try:
        n = int(value)
    except Exception:
        n = default
    return max(1, min(n, maximum))


def _v41_table_exists(table_name: str) -> bool:
    try:
        row = one("SELECT to_regclass(%s) AS name", (table_name,))
        return bool(row and row.get("name"))
    except Exception:
        return False


def _v41_columns(table_name: str) -> set[str]:
    try:
        rows = all_rows(
            """
            SELECT column_name
            FROM information_schema.columns
            WHERE table_schema = 'public' AND table_name = %s
            """,
            (table_name,),
        )
        return {str(r["column_name"]) for r in rows if r.get("column_name")}
    except Exception:
        return set()


def _v41_safe_one(sql: str, params=()):
    try:
        return one(sql, params)
    except Exception as exc:
        return {"_error": str(exc)}


def _v41_safe_all(sql: str, params=()):
    try:
        return all_rows(sql, params)
    except Exception as exc:
        return [{"_error": str(exc)}]


def _v41_proxy_json(path: str, params: dict):
    try:
        query = _v41_urlparse.urlencode(params)
        url = f"{_V41_TBM_BASE}{path}"
        if query:
            url = f"{url}?{query}"
        req = _v41_urlrequest.Request(url, headers={"Accept": "application/json"})
        with _v41_urlrequest.urlopen(req, timeout=3) as resp:
            return _v41_json.loads(resp.read().decode("utf-8"))
    except Exception:
        return None


def _v41_latest_view(device_id: str):
    # Try latest-view first, then latest.
    for path in ("/api/tbm/latest-view", "/api/tbm/latest"):
        data = _v41_proxy_json(path, {"deviceId": device_id})
        if isinstance(data, dict):
            return data.get("data") if isinstance(data.get("data"), dict) else data
    return None


def _v41_get_machine_fields(record: dict):
    if not isinstance(record, dict):
        return {}
    data = record.get("data") if isinstance(record.get("data"), dict) else record
    machine = data.get("machine") if isinstance(data.get("machine"), dict) else data
    fields = machine.get("fields") if isinstance(machine.get("fields"), dict) else data.get("fields")
    return fields if isinstance(fields, dict) else {}


def _v41_flatten_tbm_record(record: dict):
    if not isinstance(record, dict):
        return None
    data = record.get("data") if isinstance(record.get("data"), dict) else record
    fields = _v41_get_machine_fields(data)
    if not fields:
        return None
    item = {
        "deviceId": data.get("deviceId"),
        "timestamp": data.get("timestamp"),
        "receivedAt": data.get("receivedAt"),
    }
    for key in _V41_TBM_KEYS:
        f = fields.get(key)
        if isinstance(f, dict):
            item[key] = f.get("displayValue")
            item[f"{key}Status"] = f.get("status") or f.get("decodeStatus")
    return item


@app.get("/api/tbm/history")
def v41_tbm_history(deviceId: str = "DZ1360", limit: int = 300):
    n = _v41_limit(limit, 300, 2000)

    # 1) If the 19090 collector already has history, proxy it.
    proxied = _v41_proxy_json("/api/tbm/history", {"deviceId": deviceId, "limit": n})
    if isinstance(proxied, dict) and isinstance(proxied.get("data"), dict):
        items = proxied["data"].get("items")
        if isinstance(items, list) and items:
            return _v41_ok({"deviceId": deviceId, "count": len(items[-n:]), "items": items[-n:], "historySource": "tbm_receiver"}, "tbm_receiver")

    # 2) Try local JSONL generated by collector.
    items = []
    if _V41_TBM_HISTORY_JSONL.exists():
        try:
            with _V41_TBM_HISTORY_JSONL.open("r", encoding="utf-8") as f:
                lines = _v41_deque(f, maxlen=n * 5)
            for line in lines:
                try:
                    rec = _v41_json.loads(line)
                except Exception:
                    continue
                if rec.get("deviceId") != deviceId:
                    continue
                flat = _v41_flatten_tbm_record(rec)
                if flat:
                    items.append(flat)
        except Exception:
            items = []
    if items:
        return _v41_ok({"deviceId": deviceId, "count": len(items[-n:]), "items": items[-n:], "historySource": "history_jsonl"}, "history_jsonl")

    # 3) Fallback: latest only, clearly marked.
    latest = _v41_latest_view(deviceId)
    flat = _v41_flatten_tbm_record(latest or {})
    if flat:
        return _v41_ok({"deviceId": deviceId, "count": 1, "items": [flat], "historySource": "latest_only", "warning": "未找到历史序列，仅返回最新快照"}, "latest_only")

    return _v41_ok({"deviceId": deviceId, "count": 0, "items": [], "historySource": "none", "warning": "19090 latest/history unavailable"}, "none")


@app.get("/api/tbm/guidance/latest")
def v41_tbm_guidance_latest(deviceId: str = "DZ1360"):
    latest = _v41_latest_view(deviceId) or {}
    return _v41_ok({
        "deviceId": deviceId,
        "timestamp": latest.get("timestamp"),
        "receivedAt": latest.get("receivedAt"),
        "guidanceAvailable": bool(latest.get("guidance")),
        "guidanceStatus": "导向数据已接入" if latest.get("guidance") else "导向数据暂未接入",
        "guidance": latest.get("guidance"),
    }, "Data2Client")


@app.get("/api/risk-sources")
def v41_risk_sources():
    if not _v41_table_exists("risk_source"):
        return _v41_ok({"items": [], "warning": "risk_source 表不存在"}, "db")
    cols = _v41_columns("risk_source")
    point_count_sql = ""
    if _v41_table_exists("monitoring_point") and "risk_source_id" in _v41_columns("monitoring_point"):
        point_count_sql = "COUNT(mp.point_id) AS \"monitoringPointCount\""
        join_sql = "LEFT JOIN monitoring_point mp ON mp.risk_source_id = rs.risk_source_id"
        group_sql = "GROUP BY rs.risk_source_id"
    else:
        point_count_sql = "0 AS \"monitoringPointCount\""
        join_sql = ""
        group_sql = "GROUP BY rs.risk_source_id"

    def col(name, alias):
        return f'rs.{name} AS "{alias}"' if name in cols else f'NULL AS "{alias}"'

    rows = _v41_safe_all(
        f"""
        SELECT
          rs.risk_source_id::text AS "riskSourceId",
          {col("section_id", "sectionId")},
          {col("risk_name", "riskName")},
          {col("risk_type", "riskType")},
          {col("crossing_relation", "crossingRelation")},
          {col("start_mileage", "startMileage")},
          {col("end_mileage", "endMileage")},
          {col("start_mileage_m", "startMileageM")},
          {col("end_mileage_m", "endMileageM")},
          {col("protection_level", "protectionLevel")},
          {col("risk_level", "riskLevel")},
          {point_count_sql}
        FROM risk_source rs
        {join_sql}
        {group_sql}
        ORDER BY MAX(rs.start_mileage_m) NULLS LAST
        """
    )
    return _v41_ok({"items": rows}, "db")


@app.get("/api/ring-context")
def v41_ring_context(ring: int):
    if not _v41_table_exists("ring_mileage_map"):
        return _v41_ok({"ring": ring, "matched": False, "reason": "ring_mileage_map 表不存在", "riskSources": [], "events": []}, "db")

    bounds = _v41_safe_one('SELECT MIN(ring_no) AS "minRing", MAX(ring_no) AS "maxRing" FROM ring_mileage_map') or {}
    ctx = _v41_safe_one(
        """
        SELECT
          ring_id::text AS "ringId",
          section_id::text AS "sectionId",
          ring_no AS "ring",
          work_date::text AS "workDate",
          start_mileage AS "startMileage",
          end_mileage AS "endMileage",
          start_mileage_m AS "startMileageM",
          end_mileage_m AS "endMileageM",
          construction_stage AS "constructionStage",
          is_actual AS "isActual"
        FROM ring_mileage_map
        WHERE ring_no = %s
        ORDER BY is_actual DESC, work_date DESC NULLS LAST
        LIMIT 1
        """,
        (ring,),
    )
    nearest = _v41_safe_one(
        """
        SELECT
          ring_id::text AS "ringId",
          ring_no AS "ring",
          start_mileage AS "startMileage",
          end_mileage AS "endMileage",
          start_mileage_m AS "startMileageM",
          end_mileage_m AS "endMileageM",
          ABS(ring_no - %s) AS "ringDistance"
        FROM ring_mileage_map
        ORDER BY ABS(ring_no - %s)
        LIMIT 1
        """,
        (ring, ring),
    )
    if not ctx or ctx.get("_error"):
        return _v41_ok({
            "ring": ring,
            "matched": False,
            "ringContext": None,
            "riskSources": [],
            "events": [],
            "availableRange": bounds,
            "nearestRing": nearest,
            "reason": "ring_not_found_in_ring_mileage_map",
        }, "db")

    risk_sources = []
    if _v41_table_exists("risk_source") and ctx.get("sectionId") and ctx.get("startMileageM") is not None and ctx.get("endMileageM") is not None:
        risk_sources = _v41_safe_all(
            """
            SELECT
              risk_source_id::text AS "riskSourceId",
              risk_name AS "riskName",
              risk_type AS "riskType",
              crossing_relation AS "crossingRelation",
              start_mileage AS "startMileage",
              end_mileage AS "endMileage",
              start_mileage_m AS "startMileageM",
              end_mileage_m AS "endMileageM",
              protection_level AS "protectionLevel",
              risk_level AS "riskLevel"
            FROM risk_source
            WHERE section_id = %s::uuid
              AND start_mileage_m IS NOT NULL
              AND end_mileage_m IS NOT NULL
              AND %s >= start_mileage_m
              AND %s <= end_mileage_m
            ORDER BY start_mileage_m
            """,
            (ctx["sectionId"], ctx["endMileageM"], ctx["startMileageM"]),
        )

    events = []
    if _v41_table_exists("event_log"):
        events = _v41_safe_all(
            """
            SELECT
              e.event_id::text AS "eventId",
              e.event_time::text AS "eventTime",
              e.event_type AS "eventType",
              e.severity,
              e.description,
              e.possible_cause AS "possibleCause",
              e.handling_action AS "handlingAction",
              e.closure_result AS "closureResult",
              e.responsible_party AS "responsibleParty",
              e.is_shutdown AS "isShutdown",
              rs.risk_name AS "riskName"
            FROM event_log e
            LEFT JOIN risk_source rs ON e.risk_source_id = rs.risk_source_id
            WHERE e.ring_id = %s::uuid
            ORDER BY e.event_time DESC
            """,
            (ctx["ringId"],),
        )

    return _v41_ok({
        "ring": ring,
        "matched": True,
        "ringContext": ctx,
        "riskSources": risk_sources,
        "events": events,
        "availableRange": bounds,
        "nearestRing": nearest,
    }, "db")


@app.get("/api/events")
def v41_events(limit: int = 50):
    n = _v41_limit(limit, 50, 500)
    if not _v41_table_exists("event_log"):
        return _v41_ok({"items": [], "warning": "event_log 表不存在"}, "db")
    rows = _v41_safe_all(
        f"""
        SELECT
          e.event_id::text AS "eventId",
          e.event_time::text AS "eventTime",
          e.event_type AS "eventType",
          e.severity,
          e.description,
          e.possible_cause AS "possibleCause",
          e.handling_action AS "handlingAction",
          e.closure_result AS "closureResult",
          e.responsible_party AS "responsibleParty",
          e.is_shutdown AS "isShutdown",
          r.ring_no AS "ringNo",
          r.start_mileage AS "startMileage",
          r.end_mileage AS "endMileage",
          rs.risk_name AS "riskName",
          rs.risk_type AS "riskType",
          rs.risk_level AS "riskLevel"
        FROM event_log e
        LEFT JOIN ring_mileage_map r ON e.ring_id = r.ring_id
        LEFT JOIN risk_source rs ON e.risk_source_id = rs.risk_source_id
        ORDER BY e.event_time DESC
        LIMIT {n}
        """
    )
    return _v41_ok({"items": rows}, "db")


@app.get("/api/monitoring/summary")
def v41_monitoring_summary():
    if not _v41_table_exists("monitoring_reading") or not _v41_table_exists("monitoring_point"):
        return _v41_ok({"total": {}, "levelCount": [], "itemCount": [], "dateCount": [], "topAlarmPoints": []}, "db")

    total = _v41_safe_one(
        """
        SELECT
          COUNT(*) AS "totalReadingCount",
          COUNT(*) FILTER (WHERE COALESCE(alert_level::text, '') NOT IN ('', 'normal', '正常')) AS "abnormalCount",
          COUNT(DISTINCT point_id) AS "pointCount"
        FROM monitoring_reading
        """
    ) or {}

    levels = _v41_safe_all(
        """
        SELECT COALESCE(alert_level, 'unknown') AS "alertLevel", COUNT(*) AS "count"
        FROM monitoring_reading
        WHERE COALESCE(alert_level::text, '') NOT IN ('', 'normal', '正常')
        GROUP BY alert_level
        ORDER BY COUNT(*) DESC
        """
    )
    items = _v41_safe_all(
        """
        SELECT COALESCE(p.monitoring_item, '未知') AS "monitoringItem", COUNT(*) AS "count"
        FROM monitoring_reading r
        LEFT JOIN monitoring_point p ON r.point_id = p.point_id
        WHERE COALESCE(r.alert_level::text, '') NOT IN ('', 'normal', '正常')
        GROUP BY p.monitoring_item
        ORDER BY COUNT(*) DESC
        """
    )
    dates = _v41_safe_all(
        """
        SELECT r.measured_at::date::text AS "date", COUNT(*) AS "count"
        FROM monitoring_reading r
        WHERE COALESCE(r.alert_level::text, '') NOT IN ('', 'normal', '正常')
        GROUP BY r.measured_at::date
        ORDER BY r.measured_at::date
        """
    )
    top = _v41_safe_all(
        """
        WITH ranked AS (
          SELECT
            p.point_code,
            p.monitoring_item,
            r.measured_at,
            r.cumulative_change,
            r.current_value,
            r.change_rate,
            r.alert_level,
            r.source_id,
            COUNT(*) OVER (PARTITION BY p.point_code, p.monitoring_item) AS abnormal_count,
            ROW_NUMBER() OVER (PARTITION BY p.point_code, p.monitoring_item ORDER BY r.measured_at DESC) AS rn
          FROM monitoring_reading r
          LEFT JOIN monitoring_point p ON r.point_id = p.point_id
          WHERE COALESCE(r.alert_level::text, '') NOT IN ('', 'normal', '正常')
        )
        SELECT
          point_code AS "pointCode",
          monitoring_item AS "monitoringItem",
          abnormal_count AS "abnormalCount",
          measured_at::text AS "latestDate",
          cumulative_change AS "latestCumulativeChange",
          current_value AS "latestCurrentValue",
          change_rate AS "latestChangeRate",
          alert_level AS "latestAlertLevel",
          source_id::text AS "sourceId"
        FROM ranked
        WHERE rn = 1
        ORDER BY abnormal_count DESC, latestDate DESC NULLS LAST
        LIMIT 30
        """
    )

    return _v41_ok({"total": total, "levelCount": levels, "itemCount": items, "dateCount": dates, "topAlarmPoints": top}, "db")


@app.get("/api/monitoring/point-trend")
def v41_monitoring_point_trend(pointCode: str, item: Optional[str] = None, dateFrom: Optional[str] = None, dateTo: Optional[str] = None, limit: int = 500):
    n = _v41_limit(limit, 500, 5000)
    if not _v41_table_exists("monitoring_reading") or not _v41_table_exists("monitoring_point"):
        return _v41_ok({"pointCode": pointCode, "item": item, "count": 0, "items": []}, "db")

    where = ["p.point_code = %s"]
    params = [pointCode]
    if item:
        where.append("p.monitoring_item = %s")
        params.append(item)
    if dateFrom:
        where.append("r.measured_at >= %s::timestamp")
        params.append(dateFrom)
    if dateTo:
        where.append("r.measured_at <= %s::timestamp")
        params.append(dateTo)

    rows = _v41_safe_all(
        f"""
        SELECT
          r.reading_id::text AS "readingId",
          p.point_code AS "pointCode",
          p.monitoring_item AS "monitoringItem",
          p.unit,
          p.warning_threshold AS "warningThreshold",
          p.alarm_threshold AS "alarmThreshold",
          r.measured_at::text AS "measuredAt",
          r.current_value AS "currentValue",
          r.cumulative_change AS "cumulativeChange",
          r.change_rate AS "changeRate",
          r.alert_level AS "alertLevel",
          r.source_id::text AS "sourceId"
        FROM monitoring_reading r
        JOIN monitoring_point p ON r.point_id = p.point_id
        WHERE {" AND ".join(where)}
        ORDER BY r.measured_at
        LIMIT {n}
        """,
        tuple(params),
    )
    return _v41_ok({"pointCode": pointCode, "item": item, "count": len(rows), "items": rows}, "db")


@app.get("/api/evidence/by-reading")
def v41_evidence_by_reading(readingId: str):
    if not _v41_table_exists("monitoring_reading"):
        return _v41_ok({"readingId": readingId, "items": []}, "db")

    reading = _v41_safe_one(
        """
        SELECT
          r.reading_id::text AS "readingId",
          r.source_id::text AS "sourceId",
          p.point_code AS "pointCode",
          p.monitoring_item AS "monitoringItem",
          r.measured_at::text AS "measuredAt",
          r.cumulative_change AS "cumulativeChange",
          r.current_value AS "currentValue",
          r.alert_level AS "alertLevel"
        FROM monitoring_reading r
        LEFT JOIN monitoring_point p ON r.point_id = p.point_id
        WHERE r.reading_id = %s::uuid
        LIMIT 1
        """,
        (readingId,),
    )
    if not reading or reading.get("_error"):
        return _v41_ok({"readingId": readingId, "items": [], "warning": "reading not found"}, "db")

    items = []
    if _v41_table_exists("extraction_evidence"):
        cols = _v41_columns("extraction_evidence")
        # Current import design says extraction_evidence has related_table='monitoring_reading' and related_key=reading_id.
        if "related_key" in cols and "related_table" in cols:
            items = _v41_safe_all(
                """
                SELECT
                  e.evidence_id::text AS "evidenceId",
                  e.source_id::text AS "sourceId",
                  sd.file_name AS "fileName",
                  sd.file_type AS "fileType",
                  sd.document_date::text AS "documentDate",
                  e.page_no AS "pageNo",
                  e.section_title AS "sectionTitle",
                  e.table_title AS "tableTitle",
                  e.row_index AS "rowIndex",
                  e.cell_text AS "cellText",
                  e.extracted_text AS "extractedText",
                  e.confidence,
                  e.created_at::text AS "createdAt"
                FROM extraction_evidence e
                LEFT JOIN source_document sd ON e.source_id = sd.source_id
                WHERE e.related_table = 'monitoring_reading' AND e.related_key = %s
                ORDER BY e.page_no NULLS LAST, e.row_index NULLS LAST
                LIMIT 50
                """,
                (readingId,),
            )
        elif "source_id" in cols and reading.get("sourceId"):
            items = _v41_safe_all(
                """
                SELECT
                  e.evidence_id::text AS "evidenceId",
                  e.source_id::text AS "sourceId",
                  sd.file_name AS "fileName",
                  sd.file_type AS "fileType",
                  sd.document_date::text AS "documentDate",
                  e.page_no AS "pageNo",
                  NULL AS "sectionTitle",
                  NULL AS "tableTitle",
                  e.row_index AS "rowIndex",
                  NULL AS "cellText",
                  NULL AS "extractedText",
                  e.confidence,
                  e.created_at::text AS "createdAt"
                FROM extraction_evidence e
                LEFT JOIN source_document sd ON e.source_id = sd.source_id
                WHERE e.source_id = %s::uuid
                ORDER BY e.page_no NULLS LAST, e.row_index NULLS LAST
                LIMIT 50
                """,
                (reading["sourceId"],),
            )

    level = "row" if items and not items[0].get("_error") else "source_only"
    return _v41_ok({"readingId": readingId, "reading": reading, "items": items, "evidenceLevel": level}, "db")


@app.get("/api/documents/{source_id}/pages/{page_no}")
def v41_document_page(source_id: str, page_no: int):
    # Prefer staging page table if it exists.
    for table in ("stg_file_extracted_page", "stg_file_extracted_pages"):
        if _v41_table_exists(table):
            rows = _v41_safe_all(
                f"""
                SELECT *
                FROM {table}
                WHERE source_document_id::text = %s AND page_no = %s
                LIMIT 1
                """,
                (source_id, page_no),
            )
            if rows and not rows[0].get("_error"):
                return _v41_ok(rows[0], "file_staging")

    if _v41_table_exists("source_document"):
        doc = _v41_safe_one(
            """
            SELECT
              source_id::text AS "sourceId",
              file_name AS "fileName",
              file_type AS "fileType",
              document_date::text AS "documentDate",
              storage_path AS "storagePath",
              description
            FROM source_document
            WHERE source_id::text = %s
            LIMIT 1
            """,
            (source_id,),
        )
        return {"code": 404, "message": "page not found; source document found but page-level extraction is missing", "data": doc, "source": "db"}

    return {"code": 404, "message": "page not found", "data": None, "source": "db"}

# ===== V4_1_BACKEND_ENGINEERING_APIS_END =====
