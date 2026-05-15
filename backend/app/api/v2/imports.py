from fastapi import APIRouter, UploadFile, File, Form
from app.utils.field_normalizer import suggest_mapping
import csv
import io
from openpyxl import load_workbook
from docx import Document

router = APIRouter(prefix="/imports", tags=["v2-imports"])


def _preview_csv(data: bytes):
    text = data.decode("utf-8-sig", errors="ignore")
    rows = list(csv.reader(io.StringIO(text)))
    headers = [str(x).strip() for x in rows[0]] if rows else []
    sample = [dict(zip(headers, row)) for row in rows[1:6]] if headers else []
    return headers, sample


def _preview_xlsx(data: bytes):
    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True, max_row=6))
    headers = [str(x).strip() if x is not None else "" for x in rows[0]] if rows else []
    sample = [dict(zip(headers, [cell for cell in row])) for row in rows[1:6]] if headers else []
    return headers, sample


def _preview_docx(data: bytes):
    doc = Document(io.BytesIO(data))
    if not doc.tables:
        return [], []
    table = doc.tables[0]
    rows = table.rows
    headers = [cell.text.strip() for cell in rows[0].cells] if rows else []
    sample = []
    for row in rows[1:6]:
        sample.append(dict(zip(headers, [cell.text.strip() for cell in row.cells])))
    return headers, sample


@router.post("/upload")
async def upload_import(file: UploadFile = File(...), data_category: str = Form("shield_operation")):
    data = await file.read()
    name = file.filename or "uploaded"
    lower = name.lower()
    if lower.endswith(".csv"):
        headers, sample = _preview_csv(data)
    elif lower.endswith(".xlsx"):
        headers, sample = _preview_xlsx(data)
    elif lower.endswith(".docx"):
        headers, sample = _preview_docx(data)
    else:
        headers, sample = [], []
    return {
        "batchId": "preview-only-v2",
        "fileName": name,
        "dataCategory": data_category,
        "detectedHeaders": headers,
        "mappings": suggest_mapping(headers),
        "sampleRows": sample,
        "status": "preview",
    }
